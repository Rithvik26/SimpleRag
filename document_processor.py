"""
Document processor for extracting text and creating chunks
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

# Document processing libraries
import PyPDF2
import docx
from bs4 import BeautifulSoup

from extensions import ProgressTracker

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor that supports both normal and graph RAG modes."""
    
    def __init__(self, config):
        self.chunk_size = config["chunk_size"]
        self.chunk_overlap = config["chunk_overlap"]
        self.rag_mode = config.get("rag_mode", "normal")
        
        # Supported file extensions
        self.supported_extensions = {'.pdf', '.txt', '.docx', '.html', '.htm'}
        
        logger.info(f"DocumentProcessor initialized with chunk_size={self.chunk_size}, rag_mode={self.rag_mode}")
    
    def is_supported_file(self, file_path: str) -> bool:
        """Check if the file type is supported."""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions
    
    def extract_text_from_file(self, file_path: str, 
                             progress_tracker: Optional[ProgressTracker] = None) -> str:
        """Extract text from various file formats with comprehensive error handling."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {extension}. Supported formats: {self.supported_extensions}")
        
        if progress_tracker:
            progress_tracker.update(0, 100, status="extracting", 
                                   message=f"Extracting text from {path.name}",
                                   current_file=path.name)
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path, progress_tracker)
            elif extension == '.txt':
                return self._extract_from_txt(file_path, progress_tracker)
            elif extension == '.docx':
                return self._extract_from_docx(file_path, progress_tracker)
            elif extension in ['.html', '.htm']:
                return self._extract_from_html(file_path, progress_tracker)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            if progress_tracker:
                progress_tracker.update(100, 100, status="error", 
                                       message=f"Error extracting text: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str, 
                         progress_tracker: Optional[ProgressTracker] = None) -> str:
        """Extract text from PDF file with better error handling."""
        try:
            text_parts = []
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                
                if total_pages == 0:
                    logger.warning(f"PDF file has no pages: {file_path}")
                    return ""
                
                logger.debug(f"Extracting text from {total_pages} pages")
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text_parts.append(page_text)
                        
                        if progress_tracker and total_pages > 0:
                            progress_percentage = int(((i + 1) / total_pages) * 100)
                            progress_tracker.update(progress_percentage, 100, 
                                                   message=f"Extracting page {i + 1} of {total_pages}")
                    
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i + 1}: {e}")
                        continue
            
            extracted_text = "\n".join(text_parts)
            
            if not extracted_text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_txt(self, file_path: str, 
                         progress_tracker: Optional[ProgressTracker] = None) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                        text = file.read()
                        
                        if progress_tracker:
                            progress_tracker.update(100, 100, status="complete", 
                                                   message="Text file extraction complete")
                        
                        logger.info(f"Successfully extracted {len(text)} characters from TXT file using {encoding} encoding")
                        return text
                        
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with errors='ignore'
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                logger.warning(f"Used UTF-8 with error ignoring for file: {file_path}")
                return text
                
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to extract text from TXT file: {str(e)}")
    
    def _extract_from_docx(self, file_path: str, 
                          progress_tracker: Optional[ProgressTracker] = None) -> str:
        """Extract text from DOCX file with comprehensive content extraction."""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            paragraphs = doc.paragraphs
            total_paragraphs = len(paragraphs)
            
            logger.debug(f"Extracting text from {total_paragraphs} paragraphs")
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_parts.append(paragraph.text.strip())
                
                if progress_tracker and total_paragraphs > 0:
                    progress_percentage = int(((i + 1) / total_paragraphs) * 50)  # 50% for paragraphs
                    progress_tracker.update(progress_percentage, 100, 
                                           message=f"Extracting paragraph {i + 1} of {total_paragraphs}")
            
            # Extract from tables
            tables = doc.tables
            if tables:
                logger.debug(f"Extracting text from {len(tables)} tables")
                
                for table_idx, table in enumerate(tables):
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    
                    if table_text:
                        text_parts.append("\n".join(table_text))
                    
                    if progress_tracker:
                        progress_percentage = 50 + int(((table_idx + 1) / len(tables)) * 50)  # 50% for tables
                        progress_tracker.update(progress_percentage, 100, 
                                               message=f"Extracting table {table_idx + 1} of {len(tables)}")
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}")
    
    def _extract_from_html(self, file_path: str, 
                          progress_tracker: Optional[ProgressTracker] = None) -> str:
        """Extract text from HTML file with improved content extraction."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                html_content = file.read()
            
            if progress_tracker:
                progress_tracker.update(25, 100, message="Parsing HTML content")
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.extract()
            
            if progress_tracker:
                progress_tracker.update(50, 100, message="Extracting text content")
            
            # Get text content
            text = soup.get_text()
            
            if progress_tracker:
                progress_tracker.update(75, 100, message="Cleaning extracted text")
            
            # Clean up the text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            cleaned_text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if progress_tracker:
                progress_tracker.update(100, 100, message="HTML extraction complete")
            
            if not cleaned_text.strip():
                logger.warning(f"No text extracted from HTML: {file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(cleaned_text)} characters from HTML")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error reading HTML file {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to extract text from HTML file: {str(e)}")
    
    def chunk_text(self, text: str, metadata: dict, 
                   progress_tracker: Optional[ProgressTracker] = None) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks optimized for the current RAG mode."""
        if not text or not text.strip():
            logger.warning("Empty text provided for chunking")
            return []
        
        if progress_tracker:
            progress_tracker.update(0, 100, status="chunking", 
                                   message="Splitting text into chunks")
        
        # Adjust chunk size based on RAG mode
        effective_chunk_size = self.chunk_size
        if self.rag_mode == "graph":
            # Use larger chunks for graph mode to capture more context for entity extraction
            effective_chunk_size = int(self.chunk_size * 1.5)
            logger.debug(f"Using larger chunks for graph mode: {effective_chunk_size}")
        
        # Clean and normalize the text
        cleaned_text = self._clean_text(text)
        
        # Split into sentences for better chunk boundaries
        sentences = self._split_into_sentences(cleaned_text)
        total_sentences = len(sentences)
        
        if total_sentences == 0:
            logger.warning("No sentences found in text")
            return []
        
        logger.debug(f"Found {total_sentences} sentences for chunking")
        
        chunks = []
        current_chunk = ""
        current_sentence_count = 0
        
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) > effective_chunk_size and current_chunk:
                # Save current chunk
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": len(chunks),
                    "chunk_text_preview": current_chunk[:100] + "..." if len(current_chunk) > 100 else current_chunk,
                    "rag_mode": self.rag_mode,
                    "sentence_count": current_sentence_count,
                    "chunk_size": len(current_chunk)
                })
                
                chunks.append({
                    "text": current_chunk,
                    "metadata": chunk_metadata
                })
                
                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(sentences, i, current_sentence_count)
                current_chunk = " ".join(overlap_sentences) + " " + sentence if overlap_sentences else sentence
                current_sentence_count = len(overlap_sentences) + 1
            else:
                # Add sentence to current chunk
                current_chunk = potential_chunk
                current_sentence_count += 1
            
            # Update progress
            if progress_tracker and total_sentences > 0:
                progress_percentage = int(((i + 1) / total_sentences) * 100)
                progress_tracker.update(progress_percentage, 100, 
                                       message=f"Processing sentence {i + 1} of {total_sentences}")
        
        # Add the last chunk if it exists
        if current_chunk.strip():
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": len(chunks),
                "chunk_text_preview": current_chunk[:100] + "..." if len(current_chunk) > 100 else current_chunk,
                "rag_mode": self.rag_mode,
                "sentence_count": current_sentence_count,
                "chunk_size": len(current_chunk)
            })
            
            chunks.append({
                "text": current_chunk,
                "metadata": chunk_metadata
            })
        
        if progress_tracker:
            progress_tracker.update(100, 100, status="chunking_complete",
                                   message=f"Created {len(chunks)} chunks from document")
        
        logger.info(f"Created {len(chunks)} chunks from document (mode: {self.rag_mode}, avg size: {sum(len(c['text']) for c in chunks) // len(chunks) if chunks else 0})")
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for better processing."""
        # Remove excessive whitespace
        cleaned = re.sub(r'\s+', ' ', text)
        
        # Remove or replace problematic characters
        cleaned = cleaned.replace('\x00', '')  # Remove null bytes
        cleaned = cleaned.replace('\ufeff', '')  # Remove BOM
        
        # Normalize line breaks
        cleaned = cleaned.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive line breaks
        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)
        
        return cleaned.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using improved regex patterns."""
        # Pattern for sentence boundaries
        sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])|(?<=[.!?])\s*\n+\s*(?=[A-Z])'
        
        # Split by sentence boundaries
        sentences = re.split(sentence_pattern, text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Filter out very short fragments
                cleaned_sentences.append(sentence)
        
        # If regex splitting didn't work well, fall back to simple splitting
        if len(cleaned_sentences) < 2:
            # Simple fallback: split by periods followed by space and capital letter
            sentences = re.split(r'\.(?=\s+[A-Z])', text)
            cleaned_sentences = [s.strip() + '.' for s in sentences if s.strip()]
        
        return cleaned_sentences
    
    def _get_overlap_sentences(self, sentences: List[str], current_index: int, 
                              current_sentence_count: int) -> List[str]:
        """Get sentences for overlap between chunks."""
        if not sentences or current_sentence_count == 0:
            return []
        
        # Calculate number of sentences for overlap based on chunk_overlap setting
        # Estimate average sentence length and calculate overlap sentences
        avg_sentence_length = 100  # Rough estimate
        overlap_sentences_count = max(1, self.chunk_overlap // avg_sentence_length)
        overlap_sentences_count = min(overlap_sentences_count, current_sentence_count - 1)
        
        if overlap_sentences_count <= 0:
            return []
        
        # Get the last N sentences from the current chunk
        start_index = max(0, current_index - overlap_sentences_count)
        return sentences[start_index:current_index]
    
    def validate_chunks(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Validate chunks and return statistics."""
        if not chunks:
            return {
                "valid": False,
                "error": "No chunks provided",
                "total_chunks": 0
            }
        
        stats = {
            "valid": True,
            "total_chunks": len(chunks),
            "total_characters": 0,
            "average_chunk_size": 0,
            "min_chunk_size": float('inf'),
            "max_chunk_size": 0,
            "empty_chunks": 0,
            "chunks_too_large": 0,
            "chunks_too_small": 0
        }
        
        chunk_sizes = []
        
        for i, chunk in enumerate(chunks):
            # Validate chunk structure
            if not isinstance(chunk, dict):
                stats["valid"] = False
                stats["error"] = f"Chunk {i} is not a dictionary"
                return stats
            
            if "text" not in chunk:
                stats["valid"] = False
                stats["error"] = f"Chunk {i} missing 'text' field"
                return stats
            
            if "metadata" not in chunk:
                stats["valid"] = False
                stats["error"] = f"Chunk {i} missing 'metadata' field"
                return stats
            
            # Analyze chunk size
            chunk_text = chunk["text"]
            chunk_size = len(chunk_text)
            chunk_sizes.append(chunk_size)
            
            stats["total_characters"] += chunk_size
            stats["min_chunk_size"] = min(stats["min_chunk_size"], chunk_size)
            stats["max_chunk_size"] = max(stats["max_chunk_size"], chunk_size)
            
            # Count issues
            if chunk_size == 0:
                stats["empty_chunks"] += 1
            elif chunk_size < 50:  # Very small chunks
                stats["chunks_too_small"] += 1
            elif chunk_size > self.chunk_size * 2:  # Very large chunks
                stats["chunks_too_large"] += 1
        
        # Calculate averages
        if chunk_sizes:
            stats["average_chunk_size"] = sum(chunk_sizes) / len(chunk_sizes)
            stats["min_chunk_size"] = min(chunk_sizes)
        else:
            stats["min_chunk_size"] = 0
        
        # Add warnings
        warnings = []
        if stats["empty_chunks"] > 0:
            warnings.append(f"{stats['empty_chunks']} empty chunks found")
        if stats["chunks_too_small"] > 0:
            warnings.append(f"{stats['chunks_too_small']} very small chunks found")
        if stats["chunks_too_large"] > 0:
            warnings.append(f"{stats['chunks_too_large']} very large chunks found")
        
        stats["warnings"] = warnings
        
        return stats
    
    def get_supported_formats(self) -> Dict[str, str]:
        """Get information about supported file formats."""
        return {
            ".pdf": "Portable Document Format - extracts text from all pages",
            ".txt": "Plain text files - supports multiple encodings",
            ".docx": "Microsoft Word documents - extracts text and tables",
            ".html": "HTML web pages - extracts text content, removes scripts/styles",
            ".htm": "HTML web pages - same as .html"
        }
    
    def estimate_processing_time(self, file_path: str) -> Dict[str, Any]:
        """Estimate processing time for a file."""
        try:
            path = Path(file_path)
            file_size = path.stat().st_size
            extension = path.suffix.lower()
            
            # Rough estimates based on file type and size
            time_estimates = {
                ".txt": file_size / (1024 * 1024) * 0.1,  # 0.1 seconds per MB
                ".pdf": file_size / (1024 * 1024) * 2.0,   # 2 seconds per MB
                ".docx": file_size / (1024 * 1024) * 1.0,  # 1 second per MB
                ".html": file_size / (1024 * 1024) * 0.5,  # 0.5 seconds per MB
                ".htm": file_size / (1024 * 1024) * 0.5,   # 0.5 seconds per MB
            }
            
            base_time = time_estimates.get(extension, 1.0)
            
            # Add extra time for graph mode
            if self.rag_mode == "graph":
                base_time *= 3  # Graph extraction takes longer
            
            return {
                "estimated_seconds": max(1, int(base_time)),
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "complexity": "high" if self.rag_mode == "graph" else "medium"
            }
            
        except Exception as e:
            logger.error(f"Error estimating processing time: {e}")
            return {
                "estimated_seconds": 60,  # Default estimate
                "file_size_mb": 0,
                "complexity": "unknown",
                "error": str(e)
            }
    
    def get_processor_stats(self) -> Dict[str, Any]:
        """Get statistics about the document processor."""
        return {
            "rag_mode": self.rag_mode,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "supported_extensions": list(self.supported_extensions),
            "effective_chunk_size": int(self.chunk_size * 1.5) if self.rag_mode == "graph" else self.chunk_size
        }